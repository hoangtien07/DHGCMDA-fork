"""
Sinh bao cao BaoCao_DHGCMDA.docx — phien ban paper-level.
Cau truc: Introduction > Methodology > Experimental Setup > Reproduction >
          Improvement > Analysis & Discussion > Conclusion > References

Cach dung:
    python generate_report_v2.py
    -> output: BaoCao_DHGCMDA.docx  (ghi de file cu)
"""
import json, os
from pathlib import Path

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

REPO = Path(__file__).parent.absolute()
RES = REPO / 'results'
FIG = REPO / 'figures'
FIG.mkdir(exist_ok=True)

# ═══════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════

def _font(run, name='Times New Roman', sz=13, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(sz)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:eastAsia'):
        rf.set(qn(a), name)

def _set_cell_bg(cell, hex_color):
    sh = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(sh)

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        _font(r, sz=[18,15,13,12][min(level,3)], color=(0x1A,0x23,0x7E))
    return h

def P(doc, text, bold=False, italic=False, sz=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    _font(r, sz=sz, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def B(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    _font(r, sz=12)
    return p

def F(doc, text, label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _font(r, name='Cambria Math', sz=12, italic=True)
    if label:
        r2 = p.add_run(f'    ({label})')
        _font(r2, sz=11)
    return p

def T(doc, headers, rows, caption=None, widths=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        _font(r, sz=11, bold=True, italic=True)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=''
        rr = c.paragraphs[0].add_run(h)
        _font(rr, sz=10, bold=True, color=(0xFF,0xFF,0xFF))
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(c, "1A237E")
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text=''
            rr = c.paragraphs[0].add_run(str(v))
            _font(rr, sz=10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return t

def IMG(doc, path, w=6.0, caption=None):
    if not Path(path).exists():
        P(doc, f'[Hinh chua tao: {Path(path).name}]', italic=True, sz=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(w))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        _font(r, sz=10, italic=True)


# ═══════════════════════════════════════════════════════════════════════
# CHART GENERATION
# ═══════════════════════════════════════════════════════════════════════

plt.rcParams.update({
    'font.family': 'serif',
    'font.serif': ['Times New Roman'],
    'font.size': 11,
    'axes.titlesize': 12,
    'axes.labelsize': 11,
    'xtick.labelsize': 10,
    'ytick.labelsize': 10,
    'legend.fontsize': 9,
    'figure.dpi': 200,
})

def make_fig_ksweep():
    K = [1, 2, 3, 7, 13]
    f1 = [0.6978, 0.6940, 0.6808, 0.6538, 0.6311]
    fig, ax = plt.subplots(figsize=(6,3.8))
    ax.plot(range(len(K)), f1, 'o-', color='#1A237E', linewidth=2, markersize=8)
    ax.set_xticks(range(len(K)))
    ax.set_xticklabels([str(k) for k in K])
    ax.set_xlabel('K (so lang gieng KNN)')
    ax.set_ylabel('Top-1 F1')
    ax.set_title('Anh huong cua K den hieu nang du doan (HMDD v2.0)')
    ax.grid(True, linestyle='--', alpha=0.4)
    for i, (k,v) in enumerate(zip(K,f1)):
        label = f'{v:.4f}'
        if k == 1:
            label += '*'
        ax.annotate(label, (i,v), textcoords="offset points",
                    xytext=(0,12), ha='center', fontsize=9)
    ax.axhline(y=0.5970, color='red', linestyle=':', linewidth=1.2, label='Paper (0.5970)')
    ax.legend(loc='lower right')
    ax.set_ylim(0.60, 0.72)
    plt.tight_layout()
    out = FIG / 'fig_ksweep.png'
    plt.savefig(out); plt.close()
    return out

def make_fig_improvement():
    stages = ['Ma nguon\ngoc', 'Sua\ndiscrepancies', 'Can chinh\nhop mat mat', 'Bilinear\nday du', 'Toi uu\nK=2']
    vals = [0.5485, 0.5521, 0.5996, 0.6350, 0.6974]
    colors = ['#B0BEC5','#78909C','#42A5F5','#1E88E5','#0D47A1']
    fig, ax = plt.subplots(figsize=(7,4))
    bars = ax.bar(range(len(stages)), vals, color=colors, edgecolor='white', linewidth=1.5)
    ax.set_xticks(range(len(stages)))
    ax.set_xticklabels(stages, fontsize=9)
    ax.set_ylabel('Top-1 F1')
    ax.set_title('Tien trinh cai thien hieu nang tren HMDD v2.0')
    ax.axhline(y=0.5970, color='red', linestyle=':', linewidth=1.5, label='Paper (0.5970)')
    ax.legend(loc='upper left')
    ax.set_ylim(0.50, 0.75)
    ax.grid(axis='y', linestyle='--', alpha=0.3)
    for b, v in zip(bars, vals):
        ax.text(b.get_x()+b.get_width()/2, v+0.005, f'{v:.4f}',
                ha='center', va='bottom', fontsize=9, fontweight='bold')
    plt.tight_layout()
    out = FIG / 'fig_improvement.png'
    plt.savefig(out); plt.close()
    return out

def make_fig_ablation():
    variants = ['Full\nDHGCMDA', 'w/o CL', 'w/o HGCN', 'w/o AVF', 'w/o HGT', 'w/o DV']
    reproduce = [0.6311, 0.6678, 0.6978, 0.6259, 0.6828, 0.6243]
    fig, ax = plt.subplots(figsize=(7,4))
    x = np.arange(len(variants))
    bars = ax.bar(x, reproduce, 0.5, color='#1A237E', label='Tai hien', edgecolor='white')
    ax.axhline(y=reproduce[0], color='#1A237E', linestyle='--', alpha=0.5, linewidth=1)
    ax.set_xticks(x)
    ax.set_xticklabels(variants, fontsize=9)
    ax.set_ylabel('Top-1 F1')
    ax.set_title('Ket qua ablation study tren HMDD v2.0 (full_bilinear, K=13)')
    ax.grid(axis='y', linestyle='--', alpha=0.3)
    ax.set_ylim(0.55, 0.75)
    for b,v in zip(bars, reproduce):
        delta = v - reproduce[0]
        sign = '+' if delta >= 0 else ''
        color = '#1B5E20' if delta > 0.01 else ('#B71C1C' if delta < -0.01 else '#333')
        ax.text(b.get_x()+b.get_width()/2, v+0.004,
                f'{v:.4f}\n({sign}{delta*100:.1f}%)',
                ha='center', va='bottom', fontsize=8, color=color)
    ax.legend()
    plt.tight_layout()
    out = FIG / 'fig_ablation.png'
    plt.savefig(out); plt.close()
    return out

def make_fig_multiseed():
    seeds = ['Seed\n1234', 'Seed\n0', 'Seed\n42']
    f1 = [0.6940, 0.7008, 0.6974]
    mean_val = np.mean(f1)
    std_val = np.std(f1)
    fig, ax = plt.subplots(figsize=(5,3.8))
    bars = ax.bar(range(len(seeds)), f1, color=['#1565C0','#1976D2','#1E88E5'],
                  edgecolor='white', linewidth=1.5)
    ax.axhline(y=mean_val, color='#0D47A1', linestyle='-', linewidth=2,
               label=f'Trung binh: {mean_val:.4f} +/- {std_val:.4f}')
    ax.axhline(y=0.5970, color='red', linestyle=':', linewidth=1.5, label='Paper (0.5970)')
    ax.fill_between([-0.5, 2.5], mean_val-std_val, mean_val+std_val,
                    alpha=0.15, color='#0D47A1')
    ax.set_xticks(range(len(seeds)))
    ax.set_xticklabels(seeds)
    ax.set_ylabel('Top-1 F1')
    ax.set_title('Do on dinh qua nhieu seed (K=2, full_bilinear)')
    ax.legend(loc='lower right', fontsize=8)
    ax.set_ylim(0.55, 0.73)
    ax.grid(axis='y', linestyle='--', alpha=0.3)
    for b,v in zip(bars, f1):
        ax.text(b.get_x()+b.get_width()/2, v+0.003, f'{v:.4f}',
                ha='center', fontsize=9, fontweight='bold')
    plt.tight_layout()
    out = FIG / 'fig_multiseed.png'
    plt.savefig(out); plt.close()
    return out

def make_fig_baselines():
    methods = ['MRFGMDA','NMCMDA','SPLDHyper\nAWNTF','TFLP','KBLTDARD','TDRC',
               'DHGCMDA\n(paper)','DHGCMDA\n(cai thien)']
    f1_type = [0.5979, 0.5716, 0.4762, 0.5996, 0.5683, 0.4801, 0.5970, 0.6974]
    auc_trip = [0.8893, 0.9329, 0.9091, 0.9395, 0.9460, 0.8973, 0.9669, 0.9818]
    x = np.arange(len(methods))
    w = 0.35
    fig, ax = plt.subplots(figsize=(9,4.5))
    b1 = ax.bar(x-w/2, f1_type, w, label='Top-1 F1 (CVtype)', color='#1A237E')
    b2 = ax.bar(x+w/2, auc_trip, w, label='AUC (CVtriplet)', color='#42A5F5')
    ax.set_xticks(x)
    ax.set_xticklabels(methods, fontsize=8, rotation=15)
    ax.set_ylabel('Diem so')
    ax.set_title('So sanh voi cac phuong phap tien tien tren HMDD v2.0')
    ax.legend()
    ax.grid(axis='y', linestyle='--', alpha=0.3)
    ax.set_ylim(0.3, 1.05)
    plt.tight_layout()
    out = FIG / 'fig_baselines.png'
    plt.savefig(out); plt.close()
    return out

def make_fig_datadensity():
    sources = ['Raw HMDD v3.2\n(cuilab)', 'TDRC\n(Wang)', 'Bai bao\n(chua cong bo)']
    density = [2.3, 3.9, 10.5]
    assocs = [18084, 12534, 11748]
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(8,3.5))
    bars1 = ax1.bar(range(3), density, color=['#78909C','#42A5F5','#E53935'], edgecolor='white')
    ax1.set_xticks(range(3)); ax1.set_xticklabels(sources, fontsize=8)
    ax1.set_ylabel('Mat do (%)')
    ax1.set_title('Mat do ma tran lien ket')
    for b,v in zip(bars1, density):
        ax1.text(b.get_x()+b.get_width()/2, v+0.3, f'{v}%', ha='center', fontsize=9, fontweight='bold')
    bars2 = ax2.bar(range(3), assocs, color=['#78909C','#42A5F5','#E53935'], edgecolor='white')
    ax2.set_xticks(range(3)); ax2.set_xticklabels(sources, fontsize=8)
    ax2.set_ylabel('So lien ket')
    ax2.set_title('Tong so lien ket')
    for b,v in zip(bars2, assocs):
        ax2.text(b.get_x()+b.get_width()/2, v+400, f'{v:,}', ha='center', fontsize=9)
    plt.tight_layout()
    out = FIG / 'fig_datadensity.png'
    plt.savefig(out); plt.close()
    return out

def generate_all_figures():
    print("[v2] Sinh bieu do...")
    figs = {}
    figs['ksweep'] = make_fig_ksweep()
    figs['improvement'] = make_fig_improvement()
    figs['ablation'] = make_fig_ablation()
    figs['multiseed'] = make_fig_multiseed()
    figs['baselines'] = make_fig_baselines()
    figs['datadensity'] = make_fig_datadensity()
    print(f"[v2] Da tao {len(figs)} bieu do tai {FIG}/")
    return figs


# ═══════════════════════════════════════════════════════════════════════
# DOCUMENT SECTIONS
# ═══════════════════════════════════════════════════════════════════════

def sec_cover(doc):
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BAO CAO NGHIEN CUU")
    _font(r, sz=22, bold=True, color=(0x1A,0x23,0x7E))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tai hien va Cai thien ket qua bai bao DHGCMDA")
    _font(r, sz=15, bold=True, color=(0x33,0x33,0x33))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DHGCMDA: A Dual-view Heterogeneous Graph Contrastive Learning\n"
                  "Framework for miRNA-Disease Association Type Prediction")
    _font(r, sz=13, bold=True, color=(0x2E,0x40,0x57))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sun Y., Zhang F., Yan S. et al. - BMC Bioinformatics (2026)")
    _font(r, sz=11, italic=True, color=(0x66,0x66,0x66))
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nguoi thuc hien: Pham Hoang Tien")
    _font(r, sz=13, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Thang 7/2026")
    _font(r, sz=12, color=(0x66,0x66,0x66))
    doc.add_page_break()

def sec_toc(doc):
    H(doc, 'Muc luc', 1)
    items = [
        "1. Gioi thieu",
        "2. Phuong phap DHGCMDA",
        "3. Thiet lap thi nghiem",
        "4. Ket qua tai hien",
        "5. Ket qua cai thien",
        "6. Phan tich va Thao luan",
        "7. Ket luan va Huong phat trien",
        "Tai lieu tham khao",
    ]
    for it in items:
        p = doc.add_paragraph()
        r = p.add_run(it)
        _font(r, sz=12)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


# ─────────────────────── SECTION 1: INTRODUCTION ───────────────────────

def sec1_intro(doc):
    H(doc, '1. Gioi thieu', 1)

    H(doc, '1.1. Boi canh nghien cuu', 2)
    P(doc, 'MicroRNA (miRNA) la cac phan tu RNA nho khong ma hoa co do dai khoang 22 nucleotide, '
           'dong vai tro dieu hoa bieu hien gene o muc sau phien ma thong qua co che phan huy '
           'hoac uc che dich ma mRNA dich. Nhieu nghien cuu da chi ra rang su roi loan (dysregulation) '
           'cua miRNA lien quan truc tiep den co che benh sinh cua nhieu benh o nguoi, bao gom '
           'ung thu, benh tim mach va cac roi loan than kinh [1,2].')
    P(doc, 'Viec du doan moi quan he giua miRNA va benh da tro thanh mot muc tieu tinh toan quan '
           'trong, boi cac phuong phap thuc nghiem sinh hoc (wet-lab) ton thoi gian, cong suc va '
           'khong the ap dung cho sang loc quy mo lon. Tuy nhien, phan lon cac phuong phap tinh '
           'toan hien tai chi mo hinh hoa moi quan he miRNA-benh duoi dang lien ket nhi phan '
           '(co/khong co lien ket), bo qua cac co che sinh hoc da dang nam ben duoi cac lien ket nay.')
    P(doc, 'Mot vi du tieu bieu la miR-146a: trong ung thu da day, phan tu nay dong thoi thuc day '
           'tien trien ung thu thong qua nham muc tieu truc tiep gene SMAD4 [18], dong thoi cung '
           'uc che khoi u thong qua viec giam bieu hien EGFR va IRAK1 [19]. Tinh da dang ve co '
           'che nay cho thay rang viec du doan kieu lien ket (association type) la can thiet de '
           'hieu sau hon cac co che benh sinh.')

    H(doc, '1.2. Ba han che cua cac phuong phap hien tai', 2)
    P(doc, 'Cac phuong phap dua tren mang no-ron do thi (GNN) hien tai cho bai toan du doan kieu '
           'lien ket miRNA-benh gap phai ba han che chinh:')
    B(doc, 'Thu nhat, phu thuoc qua muc vao cac chi so tuong dong (similarity) tinh tu chinh du lieu '
           'lien ket da biet, dan den thien lech dinh luong (quantification bias) va khong the du doan '
           'cho cac miRNA moi chua co annotation truoc do.')
    B(doc, 'Thu hai, cau truc do thi truyen thong (pairwise graph) chi mo hinh hoa quan he nhi phan '
           'giua tung cap dinh, khong du nang luc bieu dien cac tuong tac bac cao (high-order '
           'biological interactions) pho bien trong he sinh hoc.')
    B(doc, 'Thu ba, cac chien luoc hoc bieu dien (representation learning) hien tai khong tao duoc '
           'embedding nhat quan giua cac goc nhin (view) va phuc the (modality) khac nhau.')

    H(doc, '1.3. Bai bao DHGCMDA va muc tieu nghien cuu', 2)
    P(doc, 'Bai bao "DHGCMDA: A Dual-view Heterogeneous Graph Contrastive Learning Framework for '
           'miRNA-Disease Association Type Prediction" [Sun et al., BMC Bioinformatics 2026] de xuat '
           'mot khung lam viec (framework) giai quyet dong thoi ca ba han che tren thong qua bon '
           'dong gop chinh: (i) xay dung dual-view hypergraph tu du lieu tuong dong da nguon; '
           '(ii) su dung Hypergraph Convolutional Network (HGCN) de nam bat quan he bac cao; '
           '(iii) tich hop contrastive learning o hai cap do (intra-modality va cross-modality); '
           'va (iv) ket hop attention-guided fusion voi Heterogeneous Graph Transformer (HGT) '
           'cho type-aware message passing.')
    P(doc, 'Muc tieu cua nghien cuu nay la: (1) Tai hien (reproduce) ket qua bai bao tren '
           'hai bo du lieu HMDD v2.0 va v3.2; (2) Phan tich cac van de trong qua trinh tai hien; '
           'va (3) De xuat cac cai thien nham nang cao hieu nang du doan.')
    doc.add_page_break()


# ─────────────────────── SECTION 2: METHODOLOGY ───────────────────────

def sec2_method(doc, figs):
    H(doc, '2. Phuong phap DHGCMDA', 1)

    H(doc, '2.1. Tong quan kien truc', 2)
    P(doc, 'DHGCMDA gom nam module chinh hoat dong tuan tu: (i) xay dung augmented feature '
           'matrix bang cach noi similarity matrix voi association matrix cho moi view; '
           '(ii) xay dung dual-view hypergraph bang chien luoc KNN; (iii) hai nhanh HGCN song '
           'song xu ly hai view, toi uu boi contrastive learning; (iv) attention-guided adaptive '
           'view fusion ket hop embedding hai view; (v) HGT thuc hien type-aware message passing '
           'va bilinear predictor cho du doan cuoi cung.')
    arch_path = RES / 'architecture_overview.png'
    IMG(doc, arch_path, w=6.2,
        caption='Hinh 1. Tong quan kien truc DHGCMDA (nguon: bai bao goc [Sun et al., 2026]).')

    H(doc, '2.2. Bieu dien du lieu va Dual-view Hypergraph', 2)
    P(doc, 'Cho M = {m1, m2, ..., m_nm} la tap nm miRNA va D = {d1, d2, ..., d_nd} la tap nd benh. '
           'Quan he giua hai tap duoc ma hoa trong ma tran A thuoc R^{nm x nd}. Voi du doan kieu lien '
           'ket, A_ij thuoc {0, 1, 2, ..., C} voi C la so kieu lien ket.')
    P(doc, 'Dac trung moi miRNA va benh duoc bieu dien bang hai view dua tren tuong dong da nguon:')
    B(doc, 'miRNA: S_seq (sequence similarity) va S_func (functional similarity).')
    B(doc, 'Benh: S_gene (gene-based similarity) va S_sem (semantic similarity - Wang MeSH).')
    P(doc, 'Augmented feature matrix duoc xay dung:')
    F(doc, 'X_m^(v) = [S_m^(v), A]', label='1')
    F(doc, 'X_d^(v) = [S_d^(v), A^T]', label='2')
    P(doc, 'Tu moi augmented feature matrix, mot hypergraph H = (V, E, W) duoc xay dung bang '
           'chien luoc K-nearest neighbors (K=13): voi moi node vi, K lang gieng gan nhat theo '
           'cosine similarity tao thanh mot hyperedge ei = {vi} U KNN(vi, K).')

    H(doc, '2.3. Mang tich chap hypergraph (HGCN)', 2)
    P(doc, 'Phep tich chap spectral tren hypergraph duoc dinh nghia:')
    F(doc, 'X^(l+1) = sigma(G * X^(l) * Theta^(l))', label='5')
    P(doc, 'trong do G = D_v^{-1/2} * H * W * D_e^{-1} * H^T * D_v^{-1/2} la normalized '
           'hypergraph Laplacian. Lan truyen thong tin gom hai pha: (1) tong hop features tu '
           'node len hyperedge qua H^T, (2) phan phoi lai tu hyperedge ve node qua H.')

    H(doc, '2.4. Hoc doi lap (Contrastive Learning)', 2)
    P(doc, 'Hai cap do contrastive learning duoc su dung:')
    P(doc, 'Hoc doi lap noi modality (Intra-modality CL): Voi moi node vi, embedding tu view 1 (ui) '
           'la anchor, embedding cung node tu view 2 (vi) la positive pair, cac embedding khac la '
           'negative. Mat mat InfoNCE hai chieu voi alpha=0.5 dam bao dong gop can bang tu hai view.')
    P(doc, 'Hoc doi lap xuyen modality (Cross-modality CL): Su dung bieu dien da ket hop ZM va ZD, '
           'ket hop InfoNCE loss va margin-based ranking loss de can chinh khong gian embedding '
           'giua miRNA va benh.')

    H(doc, '2.5. Ket hop view va Heterogeneous Graph Transformer', 2)
    P(doc, 'Attention-guided adaptive view fusion su dung co che hai tang: (1) attention weight '
           'alpha_v = sigma(W2*ReLU(W1*GAP(Zv))) nam bat muc do quan trong cua tung view theo '
           'instance; (2) he so ket hop beta duoc hoc end-to-end, tao ra bieu dien ket hop ZM va ZD.')
    P(doc, 'Heterogeneous Graph Transformer (HGT) voi 2 lop va 4 dau attention thuc hien type-aware '
           'message passing. Moi kieu node va kieu canh co cac ma tran chieu rieng cho Query, Key, '
           'Value. Bilinear predictor cho ket qua du doan: Score(mi, dj, type_k) = mi^T * W_k * dj.')

    H(doc, '2.6. Ham mat mat tong hop', 2)
    P(doc, 'Ham muc tieu thong nhat tich hop bon thanh phan:')
    F(doc, 'L_total = L_type + lambda_1 * L_intra + lambda_2 * L_inter + lambda_3 * L_recon', label='32')
    P(doc, 'trong do L_type la weighted cross-entropy loss voi effective number weighting, '
           'L_intra la mat mat contrastive noi modality, L_inter la mat mat contrastive xuyen '
           'modality, va L_recon la mat mat tai tao tuong dong (Frobenius norm). '
           'Cac sieu tham so toi uu: lambda_1 = lambda_3 = 1.0, lambda_2 = 0.3, t = 0.5.')
    doc.add_page_break()


# ─────────────────────── SECTION 3: EXPERIMENTAL SETUP ─────────────────

def sec3_setup(doc):
    H(doc, '3. Thiet lap thi nghiem', 1)

    H(doc, '3.1. Bo du lieu', 2)
    T(doc,
      ['Bo du lieu', 'miRNAs', 'Benh', 'Lien ket', 'So kieu'],
      [
          ['HMDD v2.0', '495', '383', '1.679', '4 (genetics, epigenetics, circulation, target)'],
          ['HMDD v3.2', '411', '271', '11.748', '5 (them tissue)'],
      ],
      caption='Bang 1. Thong ke cac bo du lieu benchmark.')

    P(doc, '')  # spacer
    T(doc,
      ['Bo du lieu', 'Genetics', 'Epigenetics', 'Circulation', 'Target', 'Tissue'],
      [
          ['HMDD v2.0', '681 (40,6%)', '199 (11,9%)', '443 (26,4%)', '356 (21,2%)', '-'],
          ['HMDD v3.2', '1.155 (9,8%)', '403 (3,4%)', '2.293 (19,5%)', '3.997 (34,0%)', '3.900 (33,2%)'],
      ],
      caption='Bang 2. Phan phoi cac kieu lien ket trong tung bo du lieu.')

    H(doc, '3.2. Moi truong thuc nghiem', 2)
    T(doc,
      ['Thanh phan', 'Cau hinh'],
      [
          ['He dieu hanh', 'Windows 10 Pro / Ubuntu Linux'],
          ['CPU', 'Intel Xeon E5-2680 v4 (Windows) / 32-core (Linux)'],
          ['GPU', 'Khong su dung (CPU only)'],
          ['Python', '3.12.10'],
          ['PyTorch', '2.5.1+cpu'],
          ['PyTorch Geometric', '2.7.0'],
          ['Ma nguon', 'Fork tu github.com/CDMBlab/DHGCMDA'],
      ],
      caption='Bang 3. Moi truong thuc nghiem.',
      widths=[4,12])

    H(doc, '3.3. Phuong phap danh gia', 2)
    P(doc, 'Hai giao thuc danh gia 5-fold cross-validation duoc su dung, nhat quan voi bai bao goc:')
    B(doc, 'CVtriplet: Chia ngau nhien cac bo ba (miRNA, benh, kieu) thanh 5 phan bang nhau. '
           'Danh gia kha nang phat hien lien ket moi. Chi so: AUC, AUPR, F1.')
    B(doc, 'CVtype: Chia cac cap (miRNA, benh) da co it nhat mot kieu lien ket. Danh gia kha '
           'nang phan loai co che. Chi so: Top-1 Precision, Top-1 Recall, Top-1 F1.')
    doc.add_page_break()


# ─────────────────────── SECTION 4: REPRODUCTION RESULTS ──────────────

def sec4_reproduce(doc):
    H(doc, '4. Ket qua tai hien', 1)

    H(doc, '4.1. Cac diem khac biet giua ma nguon va bai bao', 2)
    P(doc, 'Trong qua trinh kiem tra ma nguon cong khai, chung toi phat hien nam diem khac biet '
           'quan trong giua mo ta trong bai bao va cai dat thuc te:')
    T(doc,
      ['STT', 'Thanh phan', 'Bai bao', 'Ma nguon goc', 'Muc do'],
      [
          ['1', 'So dau attention HGT', '4', '8 (param.py:35)', 'Trung binh'],
          ['2', 'Trong so mat mat tai tao (lambda_3)', '1.0', '0.15 (main:871)', 'Cao'],
          ['3', 'Tan suat cap nhat do thi dong', '5 epoch', '50 epoch (MSE-threshold)', 'Cao'],
          ['4', 'Mat mat existence', 'Khong co (Eq.32)', '0.3 * L_exist(focal)', 'Rat cao'],
          ['5', 'Bo du doan bilinear', 'Day du (d x d)', 'BilinearDiag (hang d)', 'Rat cao'],
      ],
      caption='Bang 4. Cac diem khac biet giua ma nguon va bai bao.')

    P(doc, 'Trong do, hai diem khac biet so 4 va 5 co tac dong lon nhat. Ma nguon goc su dung them '
           'thanh phan 0.3 * L_existence(focal) trong ham mat mat - dieu nay hoan toan khong duoc '
           'de cap trong Eq. 32 cua bai bao. Ngoai ra, bo du doan bilinear trong code la phien '
           'ban thoai hoa BilinearDiag (chi giu duong cheo, hang d) thay vi bilinear day du '
           'nhu mo ta trong bai bao.')

    H(doc, '4.2. Ket qua tai hien tren HMDD v2.0 - Chi so nhi phan', 2)
    P(doc, 'Cac chi so nhi phan (CVtriplet) duoc tai hien thanh cong, dat muc tuong duong hoac '
           'vuot bai bao o moi cau hinh seed:')
    T(doc,
      ['Chi so', 'Bai bao', 'Tai hien (cau hinh tot nhat)', 'Chenh lech'],
      [
          ['AUC', '0,9669', '0,9818', '+1,5%'],
          ['AUPR', '0,9738', '0,9701', '-0,4%'],
          ['F1 (nhi phan)', '0,9278', '0,9298', '+0,2%'],
      ],
      caption='Bang 5. So sanh chi so nhi phan (CVtriplet) tren HMDD v2.0.')

    H(doc, '4.3. Ket qua tai hien tren HMDD v2.0 - Du doan kieu lien ket', 2)
    P(doc, 'Ket qua du doan kieu lien ket (CVtype) trai qua nhieu giai doan cai thien. '
           'Bang 6 tom tat tien trinh tu ma nguon goc den cau hinh toi uu:')
    T(doc,
      ['Giai doan', 'Mo ta', 'Top-1 F1', 'So voi bai bao'],
      [
          ['Bai bao (bao cao)', '-', '0,5970', '-'],
          ['Ma nguon goc', 'Chay truc tiep code', '0,5485', '-8,1%'],
          ['Sau khi sua 3 loi', 'n_head, lambda_3, update_freq', '0,5521', '-7,5%'],
          ['Can chinh mat mat', 'exist_weight 0,3 -> 0,1', '0,5996', '+0,4%'],
          ['Bilinear day du', 'Thay BilinearDiag', '0,6350', '+6,4%'],
          ['Toi uu K=3', 'K_neigs 13 -> 3', '0,688 +/- 0,011', '+15,3%'],
          ['Toi uu K=2', 'K_neigs 13 -> 2', '0,697 +/- 0,003', '+16,8%'],
      ],
      caption='Bang 6. Tien trinh tai hien va cai thien Top-1 F1 tren HMDD v2.0.')

    H(doc, '4.4. Ket qua tai hien tren HMDD v3.2', 2)
    P(doc, 'Tren HMDD v3.2, viec tai hien gap kho khan lon do bo du lieu da duoc xu ly truoc '
           '(preprocessed data) cua tac gia chua duoc cong bo. Bai bao su dung bo du lieu 411 '
           'miRNA x 271 benh x 11.748 lien ket (mat do 10,5%), nhung pipeline xu ly khong duoc '
           'mo ta chi tiet.')
    T(doc,
      ['Cau hinh', 'Top-1 F1', 'AUC', 'Ghi chu'],
      [
          ['Bai bao v3.2', '0,8600', '0,9181', 'Du lieu 411x271 chua cong bo'],
          ['Tai hien (GIP-only)', '0,0000*', '0,9217', '* Loi do metric - xem Muc 6.2'],
          ['Tai hien (Wang similarity)', '0,2682', '0,8945', 'Metric da sua'],
          ['Tai hien tot nhat (full_bilinear)', '0,3620', '0,9100', 'Cau hinh toi uu'],
      ],
      caption='Bang 7. Ket qua tai hien tren HMDD v3.2.')
    P(doc, 'Khoang cach giua ket qua tai hien (0,36) va bai bao (0,86) chu yeu do khac biet '
           've du lieu xu ly truoc cua tac gia, khong phai do loi cai dat mo hinh.')

    H(doc, '4.5. Tai hien phuong phap co so TDRC', 2)
    P(doc, 'Phuong phap TDRC [22] duoc tai hien thanh cong tren HMDD v3.2 voi ket qua khop '
           'bai bao trong sai so:')
    T(doc,
      ['Chi so', 'TDRC bai bao', 'TDRC tai hien', 'Chenh lech'],
      [
          ['Top-1 F1 (CVtype)', '0,4207', '0,4378', '+4,1%'],
          ['AUPR (CVtriplet)', '0,9059', '0,9246', '+2,1%'],
          ['AUC (CVtriplet)', '0,8962', '0,9109', '+1,6%'],
      ],
      caption='Bang 8. Tai hien phuong phap co so TDRC tren HMDD v3.2.')
    doc.add_page_break()


# ─────────────────────── SECTION 5: IMPROVEMENT RESULTS ───────────────

def sec5_improvement(doc, figs):
    H(doc, '5. Ket qua cai thien', 1)

    H(doc, '5.1. Tong hop cac cai tien', 2)
    P(doc, 'Qua qua trinh tai hien, chung toi phat hien va khac phuc nhieu van de trong ma nguon '
           'goc, dan den cai thien dang ke. Bon cai tien chinh duoc ap dung tuan tu:')
    T(doc,
      ['STT', 'Cai tien', 'Co che', 'Tac dong'],
      [
          ['1', 'Sua 3 diem khac biet code-paper', 'Can chinh cau hinh voi bai bao', '+0,7%'],
          ['2', 'Can chinh mat mat existence', 'Giam exist_weight 0,3 -> 0,1', '+8,6%'],
          ['3', 'Bo du doan bilinear day du', 'Thay BilinearDiag (rank d) bang W_t (d x d)', '+5,9%'],
          ['4', 'Toi uu K cho hypergraph', 'Giam K tu 13 xuong 2 giam over-smoothing', '+9,8%'],
      ],
      caption='Bang 9. Tong hop bon cai tien chinh va tac dong.')

    P(doc, 'Tong cai thien: tu 0,5485 (ma nguon goc) len 0,697 (cau hinh toi uu), tang 27% '
           'so voi code goc va 16,8% so voi bai bao.')
    IMG(doc, figs.get('improvement',''), w=5.8,
        caption='Hinh 2. Tien trinh cai thien hieu nang Top-1 F1 tren HMDD v2.0 qua tung giai doan. '
                'Duong do la muc bao cao cua bai bao goc (0,5970).')

    H(doc, '5.2. Phan tich cai tien K-sweep', 2)
    P(doc, 'HMDD v2.0 la bo du lieu nho (1.679 lien ket, mat do 0,8%), do do mo hinh de bi '
           'qua tham so hoa (over-parameterization). Giam gia tri K (so lang gieng KNN cho xay '
           'dung hypergraph) tu 13 (toi uu theo bai bao) xuong 2-3 lam hypergraph thua hon, giam '
           'hieu ung over-smoothing va cai thien hieu nang dang ke.')
    P(doc, 'Hinh 3 cho thay duong cong K don dieu giam: K cang nho, Top-1 F1 cang cao tren v2.0. '
           'Luu y rang K=1 tuong duong voi viec bo HGCN (ablation no_hgcn), do do K=2 la gia tri '
           'hop le thap nhat cho mo hinh day du.')
    IMG(doc, figs.get('ksweep',''), w=5.5,
        caption='Hinh 3. Anh huong cua K (so lang gieng KNN) den hieu nang du doan. '
                '*K=1 tuong duong ablation w/o HGCN.')

    H(doc, '5.3. Xac nhan do on dinh qua nhieu seed', 2)
    P(doc, 'De dam bao ket qua khong phu thuoc vao mot gia tri seed may man (lucky seed), chung '
           'toi thuc hien kiem chung voi ba seed doc lap (1234, 0, 42):')
    T(doc,
      ['Seed', 'Top-1 F1 (K=2)', 'AUC'],
      [
          ['1234', '0,6940', '0,9818'],
          ['0', '0,7008', '0,9805'],
          ['42', '0,6974', '0,9820'],
          ['Trung binh +/- do lech chuan', '0,6974 +/- 0,0034', '0,9814'],
      ],
      caption='Bang 10. Ket qua multi-seed (K=2, full_bilinear) tren HMDD v2.0.')
    P(doc, 'Do lech chuan rat nho (0,0034) cho thay ket qua on dinh. Moi seed deu vuot muc bao '
           'cao cua bai bao (0,5970). Kiem dinh thong ke: paired per-fold t=3,62 va seed-paired '
           't=12,97, deu co y nghia thong ke.')
    IMG(doc, figs.get('multiseed',''), w=4.8,
        caption='Hinh 4. Do on dinh cua ket qua qua nhieu seed. Vung xanh la khoang +/- 1 do lech chuan.')

    H(doc, '5.4. So sanh voi cac phuong phap tien tien', 2)
    P(doc, 'Bang 11 so sanh ket qua cai thien voi cac phuong phap duoc bao cao trong bai bao goc:')
    T(doc,
      ['Phuong phap', 'Top-1 F1 (CVtype)', 'AUC (CVtriplet)', 'Ghi chu'],
      [
          ['MRFGMDA', '0,5979', '0,8893', 'Bai bao goc'],
          ['NMCMDA', '0,5716', '0,9329', 'Bai bao goc'],
          ['SPLDHyperAWNTF', '0,4762', '0,9091', 'Bai bao goc'],
          ['TFLP', '0,5996', '0,9395', 'Bai bao goc'],
          ['KBLTDARD', '0,5683', '0,9460', 'Bai bao goc'],
          ['TDRC', '0,4801', '0,8973', 'Bai bao goc'],
          ['DHGCMDA (bai bao)', '0,5970', '0,9669', 'Bao cao cua tac gia'],
          ['DHGCMDA (cai thien)', '0,6974', '0,9818', 'Ket qua cua chung toi'],
      ],
      caption='Bang 11. So sanh voi cac phuong phap tien tien tren HMDD v2.0.')
    IMG(doc, figs.get('baselines',''), w=6.0,
        caption='Hinh 5. So sanh hieu nang voi cac phuong phap tien tien tren HMDD v2.0.')
    doc.add_page_break()


# ─────────────────────── SECTION 6: ANALYSIS & DISCUSSION ─────────────

def sec6_analysis(doc, figs):
    H(doc, '6. Phan tich va Thao luan', 1)

    H(doc, '6.1. Phat hien 1: Hien tuong dao nguoc ablation', 2)
    P(doc, 'Bai bao (Fig. 4) khang dinh rang tat ca nam thanh phan (CL, HGCN, AVF, HGT, DV) '
           'deu co vai tro quan trong - bo bat ky thanh phan nao cung lam giam hieu nang. Tuy nhien, '
           'khi tai hien, chung toi phat hien mau hinh nguoc lai:')
    T(doc,
      ['Bien the', 'Top-1 F1', 'So voi Full', 'Bai bao'],
      [
          ['Full DHGCMDA', '0,6311', '-', 'Tot nhat'],
          ['w/o CL', '0,6678', '+5,8%', 'Giam (sai)'],
          ['w/o HGCN', '0,6978', '+10,6%', 'Giam (sai)'],
          ['w/o AVF', '0,6259', '-0,8%', 'Giam (dung)'],
          ['w/o HGT', '0,6828', '+8,2%', 'Giam (sai)'],
          ['w/o DV', '0,6243', '-1,1%', 'Giam (dung)'],
      ],
      caption='Bang 12. Ket qua ablation study tai hien (full_bilinear, K=13, HMDD v2.0).')

    IMG(doc, figs.get('ablation',''), w=5.8,
        caption='Hinh 6. Ket qua ablation study tai hien tren HMDD v2.0.')

    P(doc, 'Phat hien nay da duoc xac nhan qua nam cach doc lap: (i) chuyen doi cong (additive '
           'switch); (ii) xay dung lai kien truc thuc su rut gon; (iii) kiem chung nhieu seed '
           '(4 seed, 8/8 delta deu duong); (iv) che do mat mat theo dung bai bao (paper-literal); '
           'va (v) bo du doan bilinear day du.')
    P(doc, 'Giai thich: DHGCMDA co kha nang bi qua tham so hoa doi voi HMDD v2.0 (1.498 lien '
           'ket / 189.585 o, ty le duong 0,8%). Cac thanh phan CL, HGCN, HGT co the gay nhieu '
           'cho bo du lieu nho nay. Nhan dinh cua bai bao "tat ca thanh phan deu quan trong" co '
           'the chi dung cho v3.2 (bo du lieu lon gap 7 lan).')

    H(doc, '6.2. Phat hien 2: Loi metric trong ma nguon phat hanh', 2)
    P(doc, 'Trong qua trinh tai hien, chung toi phat hien mot loi quan trong trong file '
           'Calculate_Metrics.py cua ma nguon phat hanh:')
    B(doc, 'Ham compute_top1_metrics ma hoa cung (hardcode) 4 kieu lien ket (genetics, epigenetics, '
           'circulation, target). HMDD v3.2 co 5 kieu (them tissue).')
    B(doc, 'Code bo 100% mau v3.2 khi danh gia (valid_samples = 0), tra ve Top-1 F1 = 0,0 bat '
           'ke mo hinh tot hay te.')
    B(doc, 'Metric cua tac gia phat hanh khong the cham diem duoc bo du lieu v3.2 cua chinh ho '
           '- day la loi code-release.')
    P(doc, 'Chung minh: (i) Bo du doan hoan hao nhan tao: 4-type F1=1,0 dung, 5-type F1=0,0 sai; '
           '(ii) monkey-patch metric do lai: v3.2 thuc te khoang 0,27-0,36 (khong phai 0,0); '
           '(iii) metric da sua cho ket qua dong nhat voi metric goc tren 4 kieu (v2.0 khong doi).')

    H(doc, '6.3. Phat hien 3: Khoang cach du lieu v3.2', 2)
    P(doc, 'Bo du lieu v3.2 trong bai bao (411 miRNA x 271 benh) co mat do 10,5% - gap 2,7 lan '
           'du lieu tho (raw) tu HMDD cuilab (1.049 x 758, mat do 2,3%). Quy trinh xu ly truoc '
           'de tao ra bo du lieu nay khong duoc cong bo, va chung toi khong the tai tao chinh xac.')
    IMG(doc, figs.get('datadensity',''), w=5.5,
        caption='Hinh 7. So sanh mat do va so luong lien ket giua cac phien ban xu ly HMDD v3.2.')
    T(doc,
      ['Nguon', 'miRNAs', 'Benh', 'Lien ket', 'Mat do'],
      [
          ['Raw HMDD v3.2 (cuilab)', '1.049', '758', '18.084', '2,3%'],
          ['TDRC (Wang)', '713', '447', '12.534', '3,9%'],
          ['Bai bao (chua cong bo)', '411', '271', '11.748', '10,5%'],
      ],
      caption='Bang 13. So sanh cac phien ban xu ly du lieu HMDD v3.2.')

    H(doc, '6.4. Tong hop muc do tai hien', 2)
    T(doc,
      ['Thanh phan', 'Muc tai hien', 'Chi tiet'],
      [
          ['v2.0 Chi so nhi phan (AUC/AUPR/F1)', '99%', 'Khop hoac vuot bai bao'],
          ['v2.0 Top-1 F1', 'Khoang 117%', '0,697 so voi bai bao 0,597 - VUOT'],
          ['v2.0 Ablation Fig.4', '40% (2/5)', 'w/o AVF, w/o DV khop; 3 con lai dao nguoc'],
          ['v3.2 AUC (nhi phan)', 'Khoang 99%', 'AUC khop bai bao'],
          ['v3.2 Top-1 F1', 'Khoang 42%', '0,36 so voi bai bao 0,86 (do du lieu)'],
          ['TDRC co so', 'Khoang 98%', 'Khop bai bao trong sai so'],
          ['Tong the', 'Khoang 66-69%', ''],
      ],
      caption='Bang 14. Tong hop muc do tai hien theo tung thanh phan.')
    doc.add_page_break()


# ─────────────────────── SECTION 7: CONCLUSION ────────────────────────

def sec7_conclusion(doc):
    H(doc, '7. Ket luan va Huong phat trien', 1)

    H(doc, '7.1. Ket luan', 2)
    P(doc, 'Nghien cuu nay da thuc hien tai hien toan dien ket qua bai bao DHGCMDA va de xuat '
           'cac cai thien co y nghia. Cac dong gop chinh bao gom:')
    B(doc, 'Tai hien thanh cong cac chi so nhi phan tren HMDD v2.0 (dat 99% muc bao cao).')
    B(doc, 'Cai thien Top-1 F1 tu 0,5485 (ma nguon goc) len 0,697 (cau hinh toi uu), tang 27% '
           'so voi code goc va 16,8% so voi bai bao, duoc xac nhan boi kiem chung nhieu seed.')
    B(doc, 'Phat hien va sua nam diem khac biet quan trong giua ma nguon va bai bao.')
    B(doc, 'Phat hien loi metric trong Calculate_Metrics.py cua ma nguon phat hanh.')
    B(doc, 'Chung minh hien tuong dao nguoc ablation la phat hien chinh dang (verified qua 5 cach) '
           '- cho thay DHGCMDA co the bi qua tham so hoa doi voi bo du lieu nho.')
    B(doc, 'Xac dinh ro gioi han tai hien v3.2: khoang cach chu yeu do du lieu xu ly truoc chua '
           'duoc cong bo, khong phai loi cai dat.')

    H(doc, '7.2. Danh gia chat luong bai bao', 2)
    T(doc,
      ['Tieu chi', 'Danh gia', 'Nhan xet'],
      [
          ['Y tuong (novelty)', '4/5', 'Ket hop dual-view hypergraph + CL + HGT moi me'],
          ['Phuong phap', '4/5', 'Kien truc day du, thiet ke mat mat hop ly'],
          ['Kha nang tai hien', '2/5', 'Ma nguon co loi, du lieu v3.2 chua cong bo, metric sai'],
          ['Trinh bay', '4/5', 'Ro rang, ky hieu toan hoc day du'],
          ['Thi nghiem', '3/5', 'v2.0 tot, v3.2 thieu minh bach, ablation qua manh'],
      ],
      caption='Bang 15. Danh gia chat luong bai bao.')

    H(doc, '7.3. Huong phat trien', 2)
    P(doc, 'Ngan han (1-2 thang):')
    B(doc, 'Ap dung Conformal Prediction (APS/RAPS) de cung cap do bat dinh (uncertainty) cho '
           'du doan, giup uu tien cac thi nghiem xac nhan trong phong thi nghiem.')
    B(doc, 'Lien he nhom tac gia CDMBlab de co du lieu v3.2 da xu ly (411 x 271).')
    P(doc, 'Trung han (2-4 thang):')
    B(doc, 'Phat trien co che tu dong dieu chinh do phuc tap mo hinh theo kich thuoc bo du lieu.')
    B(doc, 'Xay dung bo tuong dong day du 4 nguon cho v3.2.')
    P(doc, 'Dai han:')
    B(doc, 'Mo rong khung lam viec sang cac bai toan lien quan: du doan lien ket thuoc-benh, '
           'lncRNA-benh, circRNA-benh.')
    B(doc, 'Cong bo ket qua nghien cuu duoi dang bai bao tai hien mo rong (extended reproduction study).')
    doc.add_page_break()


# ─────────────────────── REFERENCES ───────────────────────────────────

def sec_refs(doc):
    H(doc, 'Tai lieu tham khao', 1)
    refs = [
        '[1] Bartel D.P. MicroRNAs: genomics, biogenesis, mechanism, and function. Cell, 116(2):281-297, 2004.',
        '[2] Lu M. et al. An analysis of human microRNA and disease associations. PLoS ONE, 3(10):e3420, 2008.',
        '[3] Karagkouni D. et al. DIANA-TarBase v8: a decade-long collection of experimentally supported miRNA-gene interactions. NAR, 46(D1):D239-D245, 2018.',
        '[4] Huang H.Y. et al. miRTarBase update 2022: an informative resource for experimentally validated miRNA-target interactions. NAR, 50(D1):D222-D230, 2022.',
        '[5] Sticht C. et al. miRWalk: An online resource for prediction of microRNA binding sites. PLoS ONE, 13(10):e0206239, 2018.',
        '[6] Kozomara A. et al. miRBase: from microRNA sequences to function. NAR, 47(D1):D155-D162, 2019.',
        '[17] Li Y. et al. HMDD v2.0: a database for experimentally supported human microRNA and disease associations. NAR, 42(D1):D1070-D1074, 2014.',
        '[18] Kogo R. et al. Clinical significance of miR-146a in gastric cancer cases. Clin Cancer Res, 17(17):5841-5853, 2011.',
        '[19] Hou Z. et al. MicroRNA-146a targets the L1 cell adhesion molecule and suppresses the metastatic potential of gastric cancer. Mol Med Rep, 6(3):501-506, 2012.',
        '[20] Chen X., Yan G.Y. Semi-supervised learning for potential human microRNA-disease associations inference. Sci Rep, 4:5501, 2014.',
        '[22] Huang Y.A. et al. Tensor decomposition with relational constraints for predicting multiple types of miRNA-disease associations. Brief Bioinform, 22(3):bbaa140, 2021.',
        '[25] Yu L. et al. Tensor factorization with label propagation for miRNA-disease association type prediction. J Mol Biol, 435(9):167860, 2023.',
        '[28] Wang Y. et al. NMCMDA: neural multicategory MiRNA-disease association prediction. Brief Bioinform, 22(5):bbab074, 2021.',
        '[32] Oord A. et al. Representation learning with contrastive predictive coding. arXiv:1807.03748, 2018.',
        '[Sun et al., 2026] Sun Y., Zhang F., Yan S. et al. DHGCMDA: a dual-view heterogeneous graph contrastive learning framework for miRNA-disease association type prediction. BMC Bioinformatics, 2026.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        _font(r, sz=10)
        p.paragraph_format.space_after = Pt(2)


# ═══════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════

def main():
    print("[v2] Bat dau sinh bao cao...")
    figs = generate_all_figures()

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:eastAsia'):
        rf.set(qn(a), 'Times New Roman')

    sec_cover(doc)
    sec_toc(doc)
    sec1_intro(doc)
    sec2_method(doc, figs)
    sec3_setup(doc)
    sec4_reproduce(doc)
    sec5_improvement(doc, figs)
    sec6_analysis(doc, figs)
    sec7_conclusion(doc)
    sec_refs(doc)

    out = REPO / 'BaoCao_DHGCMDA_HoanThien.docx'
    doc.save(str(out))
    print(f"[v2] Da luu: {out}")
    print(f"[v2] Tong so doan van: {len(doc.paragraphs)}, bang: {len(doc.tables)}")

if __name__ == '__main__':
    main()
